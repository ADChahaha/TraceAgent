from __future__ import annotations

from dataclasses import fields
from io import BytesIO

import pytest
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.utils import ImageReader
from ocr_processor import FileType, ProcessResult, process
from ocr_processor.impl.doc import docling_adapter as doc_docling_adapter
from ocr_processor.impl.pdf import docling_adapter as pdf_docling_adapter
from ocr_processor.schemas import BoundingBox, ContentBlock
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document


class DummyUploadFile:
    def __init__(self, filename: str, content: bytes, content_type: str | None = None):
        self.filename = filename
        self.content_type = content_type
        self.file = BytesIO(content)


def build_pdf_bytes(text: str) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.drawString(72, 720, text)
    pdf.save()
    return buffer.getvalue()


def build_scanned_pdf_bytes(text: str) -> bytes:
    font_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    try:
        font = ImageFont.truetype(font_path, 120)
    except OSError:
        pytest.skip(f"Required OCR test font was not found: {font_path}")

    image = Image.new("RGB", (1400, 500), "white")
    draw = ImageDraw.Draw(image)
    draw.text((100, 140), text, fill="black", font=font)

    image_buffer = BytesIO()
    image.save(image_buffer, format="PNG")
    image_buffer.seek(0)

    pdf_buffer = BytesIO()
    pdf = canvas.Canvas(pdf_buffer, pagesize=letter)
    pdf.drawImage(ImageReader(image_buffer), 36, 400, width=540, height=190)
    pdf.save()
    return pdf_buffer.getvalue()


def build_docx_bytes(paragraphs: list[str]) -> bytes:
    buffer = BytesIO()
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(buffer)
    return buffer.getvalue()


def build_docx_with_table_bytes(*, paragraphs: list[str], table_rows: list[list[str]]) -> bytes:
    buffer = BytesIO()
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        for row_values in table_rows:
            row_cells = table.add_row().cells
            for index, value in enumerate(row_values):
                row_cells[index].text = value

    document.save(buffer)
    return buffer.getvalue()


def build_structured_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.add_run("Project Proposal")

    document.add_paragraph("")
    document.add_paragraph("Overview")
    document.add_paragraph("")
    document.add_paragraph("This is the first body paragraph.")

    document.save(buffer)
    return buffer.getvalue()


def build_repeated_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("N/A")
    document.add_paragraph("N/A")
    document.save(buffer)
    return buffer.getvalue()


class FakeDoclingTextItem:
    def __init__(self, text: str):
        self.text = text
        self.prov = []


class FakeDoclingDocument:
    def __init__(self, texts):
        self.texts = texts


class FakeDoclingConversionResult:
    def __init__(self, texts):
        self.document = FakeDoclingDocument(texts)


def test_process_result_exposes_minimal_public_fields():
    assert [field.name for field in fields(ProcessResult)] == [
        "file_type",
        "filename",
        "md_list",
        "markdown",
        "blocks",
        "meta_info",
        "warnings",
    ]


def test_process_infers_pdf_type_from_filename(monkeypatch):
    pdf_block = ContentBlock(
        text="Hello PDF World",
        page_no=1,
        bbox=BoundingBox(x0=72.0, y0=60.0, x1=180.0, y1=72.0),
        kind="text",
        meta_info={},
    )
    monkeypatch.setattr(
        pdf_docling_adapter,
        "convert_pdf_with_docling",
        lambda content, filename: object(),
    )
    monkeypatch.setattr(
        pdf_docling_adapter,
        "build_blocks_from_docling_result",
        lambda conversion_result, pdf_bytes=None: [pdf_block],
    )

    pdf_bytes = build_pdf_bytes("Hello PDF World")
    file_obj = DummyUploadFile(
        filename="sample.pdf",
        content=pdf_bytes,
        content_type="application/pdf",
    )

    result = process(file_obj)

    assert isinstance(result, ProcessResult)
    assert result.file_type == FileType.PDF
    assert result.filename == "sample.pdf"
    assert result.blocks
    assert result.md_list
    assert result.md_list[0] == result.blocks[0].meta_info["md"]
    assert result.blocks[0].page_no == 1
    assert result.blocks[0].bbox is not None
    assert "Hello" in result.blocks[0].text
    assert "Hello PDF World" in result.markdown
    assert result.meta_info["block_count"] == len(result.blocks)
    assert result.meta_info["engine"] == "docling_rapidocr"
    assert hasattr(result, "warnings")


def test_process_pdf_raises_when_docling_conversion_fails(monkeypatch):
    def raise_docling_failure(content, filename):
        raise RuntimeError("docling pdf conversion failed")

    monkeypatch.setattr(pdf_docling_adapter, "convert_pdf_with_docling", raise_docling_failure)

    pdf_bytes = build_pdf_bytes("Hello PDF World")
    file_obj = DummyUploadFile(
        filename="sample.pdf",
        content=pdf_bytes,
        content_type="application/pdf",
    )

    with pytest.raises(RuntimeError, match="docling pdf conversion failed"):
        process(file_obj)


def test_process_infers_docx_type_from_filename(monkeypatch):
    monkeypatch.setattr(
        doc_docling_adapter,
        "convert_with_docling",
        lambda content, filename: FakeDoclingConversionResult([FakeDoclingTextItem("Hello DOCX World")]),
    )

    docx_bytes = build_docx_bytes(["First paragraph", "Second paragraph"])
    file_obj = DummyUploadFile(
        filename="sample.docx",
        content=docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    result = process(file_obj)

    assert isinstance(result, ProcessResult)
    assert result.file_type == FileType.DOCX
    assert result.filename == "sample.docx"
    assert result.blocks
    assert result.md_list
    assert result.blocks[0].kind == "text"
    assert result.blocks[0].page_no is None
    assert result.blocks[0].bbox is None
    assert result.blocks[0].text == "Hello DOCX World"
    assert result.md_list == ["Hello DOCX World"]
    assert result.markdown == "Hello DOCX World"
    assert result.meta_info["block_count"] == len(result.blocks)


def test_process_falls_back_when_docling_docx_conversion_fails(monkeypatch):
    def raise_docling_failure(content, filename):
        raise RuntimeError("docling docx conversion failed")

    monkeypatch.setattr(doc_docling_adapter, "convert_with_docling", raise_docling_failure)

    docx_bytes = build_docx_bytes(["First paragraph", "Second paragraph"])
    file_obj = DummyUploadFile(
        filename="sample.docx",
        content=docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    result = process(file_obj)

    assert isinstance(result, ProcessResult)
    assert result.file_type == FileType.DOCX
    assert result.filename == "sample.docx"
    assert result.blocks
    assert [block.text for block in result.blocks] == ["First paragraph", "Second paragraph"]
    assert all(block.page_no is None for block in result.blocks)
    assert all(block.bbox is None for block in result.blocks)
    assert result.md_list == ["First paragraph", "Second paragraph"]
    assert result.markdown == "First paragraph\n\nSecond paragraph"
    assert result.meta_info["fallback_used"] is True
    assert result.warnings == [
        "Docling DOCX pipeline failed; used python-docx fallback.",
        "docling docx conversion failed",
    ]


def test_process_falls_back_when_docling_docx_returns_no_blocks(monkeypatch):
    monkeypatch.setattr(
        doc_docling_adapter,
        "convert_with_docling",
        lambda content, filename: FakeDoclingConversionResult([]),
    )

    docx_bytes = build_docx_bytes(["First paragraph", "Second paragraph"])
    file_obj = DummyUploadFile(
        filename="sample.docx",
        content=docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    result = process(file_obj)

    assert isinstance(result, ProcessResult)
    assert result.file_type == FileType.DOCX
    assert result.filename == "sample.docx"
    assert [block.text for block in result.blocks] == ["First paragraph", "Second paragraph"]
    assert all(block.page_no is None for block in result.blocks)
    assert all(block.bbox is None for block in result.blocks)
    assert result.md_list == ["First paragraph", "Second paragraph"]
    assert result.markdown == "First paragraph\n\nSecond paragraph"
    assert result.meta_info["fallback_used"] is True
    assert result.warnings == [
        "Docling DOCX pipeline returned no blocks; used python-docx fallback."
    ]


def test_process_docx_fallback_preserves_table_as_markdown(monkeypatch):
    monkeypatch.setattr(
        doc_docling_adapter,
        "convert_with_docling",
        lambda content, filename: FakeDoclingConversionResult([]),
    )

    docx_bytes = build_docx_with_table_bytes(
        paragraphs=["实验报告"],
        table_rows=[
            ["题目", ""],
            ["姓名", "张三"],
        ],
    )
    file_obj = DummyUploadFile(
        filename="sample.docx",
        content=docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    result = process(file_obj)

    assert result.file_type == FileType.DOCX
    assert "实验报告" in result.markdown
    assert any(item.startswith("# ") for item in result.md_list)
    assert "| 题目 |  |" in result.markdown
    assert "| 姓名 | 张三 |" in result.markdown


def test_process_docx_fallback_detects_generic_headings(monkeypatch):
    monkeypatch.setattr(
        doc_docling_adapter,
        "convert_with_docling",
        lambda content, filename: FakeDoclingConversionResult([]),
    )

    docx_bytes = build_structured_docx_bytes()
    file_obj = DummyUploadFile(
        filename="structured.docx",
        content=docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    result = process(file_obj)

    assert result.file_type == FileType.DOCX
    assert "# Project Proposal" in result.markdown
    assert "# Overview" in result.markdown
    assert "This is the first body paragraph." in result.markdown


def test_process_docx_fallback_preserves_adjacent_repeated_blocks(monkeypatch):
    monkeypatch.setattr(
        doc_docling_adapter,
        "convert_with_docling",
        lambda content, filename: FakeDoclingConversionResult([]),
    )

    docx_bytes = build_repeated_docx_bytes()
    file_obj = DummyUploadFile(
        filename="repeated.docx",
        content=docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    result = process(file_obj)

    assert result.file_type == FileType.DOCX
    assert [block.text for block in result.blocks] == ["N/A", "N/A"]
    assert result.md_list == ["N/A", "N/A"]
    assert result.markdown == "N/A\n\nN/A"


def test_process_reports_legacy_doc_as_unimplemented():
    doc_bytes = b"legacy-doc-placeholder"
    file_obj = DummyUploadFile(
        filename="sample.doc",
        content=doc_bytes,
        content_type="application/msword",
    )

    result = process(file_obj)

    assert isinstance(result, ProcessResult)
    assert result.file_type == FileType.DOC
    assert result.filename == "sample.doc"
    assert result.blocks == []
    assert result.md_list == []
    assert result.markdown == ""
    assert result.meta_info["block_count"] == 0
    assert result.warnings == ["Legacy .doc processing is not implemented yet."]


def test_process_allows_explicit_type_override(monkeypatch):
    monkeypatch.setattr(
        pdf_docling_adapter,
        "convert_pdf_with_docling",
        lambda content, filename: object(),
    )
    monkeypatch.setattr(
        pdf_docling_adapter,
        "build_blocks_from_docling_result",
        lambda conversion_result, pdf_bytes=None: [
            ContentBlock(
                text="Explicit PDF",
                page_no=1,
                bbox=BoundingBox(x0=72.0, y0=60.0, x1=160.0, y1=72.0),
                kind="text",
                meta_info={},
            )
        ],
    )

    pdf_bytes = build_pdf_bytes("Explicit PDF")
    file_obj = DummyUploadFile(
        filename="unknown.bin",
        content=pdf_bytes,
    )

    result = process(file_obj, "pdf")

    assert result.file_type == FileType.PDF
    assert result.blocks
    assert result.md_list


def test_process_extracts_bbox_from_scanned_pdf():
    artifacts_path = pdf_docling_adapter.resolve_docling_artifacts_path()
    if artifacts_path is None or not artifacts_path.exists():
        pytest.skip("Docling PDF artifacts are not available for scanned PDF testing.")

    pdf_bytes = build_scanned_pdf_bytes("Hello Scan PDF")
    file_obj = DummyUploadFile(
        filename="scan.pdf",
        content=pdf_bytes,
        content_type="application/pdf",
    )

    result = process(file_obj)

    assert result.file_type == FileType.PDF
    assert result.blocks
    assert any("Hello" in block.text for block in result.blocks)

    assert result.blocks[0].page_no == 1

    bbox_blocks = [block for block in result.blocks if block.bbox is not None]
    if bbox_blocks:
        first_bbox_block = bbox_blocks[0]
        assert first_bbox_block.bbox.x1 > first_bbox_block.bbox.x0
        assert first_bbox_block.bbox.y1 > first_bbox_block.bbox.y0
