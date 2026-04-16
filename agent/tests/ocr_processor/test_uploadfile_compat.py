from __future__ import annotations

from io import BytesIO
import importlib.util
from pathlib import Path
import sys
import types

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))


def _install_docling_stub() -> None:
    if importlib.util.find_spec("docling") is not None or "docling" in sys.modules:
        return

    docling_module = types.ModuleType("docling")
    docling_module.__path__ = []

    backend_module = types.ModuleType("docling.backend")
    backend_module.__path__ = []
    pypdfium_backend_module = types.ModuleType("docling.backend.pypdfium2_backend")

    datamodel_module = types.ModuleType("docling.datamodel")
    datamodel_module.__path__ = []
    base_models_module = types.ModuleType("docling.datamodel.base_models")
    pipeline_options_module = types.ModuleType("docling.datamodel.pipeline_options")

    document_converter_module = types.ModuleType("docling.document_converter")

    class DocumentStream:
        def __init__(self, name: str, stream):
            self.name = name
            self.stream = stream

    class InputFormat:
        PDF = "pdf"

    class PdfPipelineOptions:
        def __init__(self, *args, **kwargs):
            self.args = args
            self.kwargs = kwargs

    class RapidOcrOptions:
        def __init__(self, *args, **kwargs):
            self.args = args
            self.kwargs = kwargs

    class DocumentConverter:
        def __init__(self, *args, **kwargs):
            self.args = args
            self.kwargs = kwargs

        def convert(self, *args, **kwargs):
            raise RuntimeError("docling stub should be monkeypatched in tests")

    class PdfFormatOption:
        def __init__(self, *args, **kwargs):
            self.args = args
            self.kwargs = kwargs

    class PyPdfiumDocumentBackend:
        pass

    base_models_module.DocumentStream = DocumentStream
    base_models_module.InputFormat = InputFormat
    pipeline_options_module.PdfPipelineOptions = PdfPipelineOptions
    pipeline_options_module.RapidOcrOptions = RapidOcrOptions
    document_converter_module.DocumentConverter = DocumentConverter
    document_converter_module.PdfFormatOption = PdfFormatOption
    pypdfium_backend_module.PyPdfiumDocumentBackend = PyPdfiumDocumentBackend

    sys.modules["docling"] = docling_module
    sys.modules["docling.backend"] = backend_module
    sys.modules["docling.backend.pypdfium2_backend"] = pypdfium_backend_module
    sys.modules["docling.datamodel"] = datamodel_module
    sys.modules["docling.datamodel.base_models"] = base_models_module
    sys.modules["docling.datamodel.pipeline_options"] = pipeline_options_module
    sys.modules["docling.document_converter"] = document_converter_module


_install_docling_stub()

from docx import Document
from starlette.datastructures import UploadFile as StarletteUploadFile

from ocr_processor import FileType, process
from ocr_processor.impl.doc import docling_adapter as doc_docling_adapter


class FakeDoclingTextItem:
    def __init__(self, text: str):
        self.text = text
        self.prov = []


class FakeDoclingDocument:
    def __init__(self, texts):
        self.texts = texts


class FakeDoclingConversionResult:
    def __init__(self, texts):
        self.document = FakeDoclingDocument(texts)


def _build_docx_bytes(paragraphs: list[str]) -> bytes:
    buffer = BytesIO()
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(buffer)
    return buffer.getvalue()


def test_process_accepts_starlette_upload_file(monkeypatch):
    monkeypatch.setattr(
        doc_docling_adapter,
        "convert_with_docling",
        lambda content, filename: FakeDoclingConversionResult([FakeDoclingTextItem("Hello UploadFile")]),
    )

    upload_file = StarletteUploadFile(
        file=BytesIO(_build_docx_bytes(["First paragraph"])),
        filename="upload.docx",
    )

    result = process(upload_file)

    assert result.file_type == FileType.DOCX
    assert result.filename == "upload.docx"
    assert result.md_list == ["Hello UploadFile"]
    assert result.markdown == "Hello UploadFile"
