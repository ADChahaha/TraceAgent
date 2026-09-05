"""真实 DOCX 经统一准备接口生成资源，再通过路径发起问答。"""

import io
from pathlib import Path

import numpy as np
import pytest
from docx import Document
from fastapi.testclient import TestClient

from main import create_app
from service.document_resources import model as embedding_model


@pytest.fixture
def resources(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCUMENT_RESOURCES_ROOT", str(tmp_path))
    calls = []

    class Embedder:
        def encode(self, texts):
            calls.extend(texts)
            return np.array([[1.0, 0.0] for _ in texts], dtype=np.float32)

    monkeypatch.setattr(embedding_model, "get_embedder", lambda **kwargs: Embedder())
    monkeypatch.setattr(embedding_model, "get_tokenizer", lambda *a, **k: lambda text: [(i, i + 1) for i in range(len(text))])
    return tmp_path, calls


def upload(client):
    document = Document()
    document.add_heading("合同", 1)
    document.add_paragraph("付款期限为三十天。")
    data = io.BytesIO()
    document.save(data)
    return client.post("/v1/document-resources", files=[
        ("files", ("合同.docx", data.getvalue(), "application/octet-stream")),
        ("files", ("附件.docx", data.getvalue(), "application/octet-stream")),
    ])


def test_prepare_real_docx_publishes_complete_resource(resources):
    root, calls = resources
    response = upload(TestClient(create_app()))
    assert response.status_code == 200, response.text
    result = response.json()
    path = Path(result["resource_path"])
    assert path.parent == root
    assert (path / "manifest.json").is_file()
    assert (path / "index" / "vectors.npy").is_file()
    assert len(list((path / "documents").iterdir())) == 2
    assert [doc["filename"] for doc in result["documents"]] == ["合同.docx", "附件.docx"]
    assert "三十天" in result["documents"][0]["html"]
    assert calls


def test_prepare_failure_does_not_publish_resource(resources, monkeypatch):
    root, _ = resources

    def fail(**kwargs):
        raise RuntimeError("embedding unavailable")

    monkeypatch.setattr(embedding_model, "get_embedder", fail)
    response = upload(TestClient(create_app(), raise_server_exceptions=False))
    assert response.status_code == 500
    assert "embedding unavailable" in response.text
    assert list(root.iterdir()) == []


def test_prepare_rejects_unsupported_file(resources):
    response = TestClient(create_app()).post("/v1/document-resources", files={"files": ("bad.txt", b"text")})
    assert response.status_code == 422
    assert list(resources[0].iterdir()) == []


def test_qa_uses_prepared_path_without_rebuilding_or_deleting(resources, monkeypatch):
    from langchain_core.messages import AIMessage
    from service.file_extraction_agent import manager

    class Model:
        def bind_tools(self, tools):
            return self

        def invoke(self, messages):
            return AIMessage(content="回答", response_metadata={"finish_reason": "stop"})

    monkeypatch.setattr(manager, "build_qa_model", lambda config: Model())
    client = TestClient(create_app())
    prepared = upload(client)
    assert prepared.status_code == 200, prepared.text
    path = prepared.json()["resource_path"]
    before = list(resources[1])
    for cid in ("cmp_first", "cmp_second"):
        response = client.post("/v1/document-qa/chat/completions", json={
            "completion_id": cid, "resource_path": path,
            "messages": [{"role": "user", "content": "你好"}],
        })
        assert response.status_code == 200, response.text
        assert "completion.completed" in response.text
        assert Path(path).is_dir()
    assert resources[1] == before


def test_qa_rejects_unmanaged_resource_path(resources):
    response = TestClient(create_app()).post("/v1/document-qa/chat/completions", json={
        "completion_id": "cmp_invalid", "resource_path": str(resources[0].parent),
        "messages": [{"role": "user", "content": "你好"}],
    })
    assert response.status_code == 422


@pytest.mark.parametrize("damage", ["missing_index", "bad_version", "outside_reference"])
def test_qa_rejects_damaged_resource_without_rebuilding(resources, damage):
    import json

    client = TestClient(create_app())
    prepared = upload(client)
    assert prepared.status_code == 200, prepared.text
    path = Path(prepared.json()["resource_path"])
    if damage == "missing_index":
        (path / "index" / "vectors.npy").unlink()
    elif damage == "bad_version":
        manifest_path = path / "manifest.json"
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        manifest["version"] = -1
        manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
    else:
        index_path = path / "index" / "index.json"
        index = json.loads(index_path.read_text(encoding="utf-8"))
        index["chunks"][0]["covered_files"] = ["../../outside.md"]
        index_path.write_text(json.dumps(index), encoding="utf-8")
    before = list(resources[1])
    response = client.post("/v1/document-qa/chat/completions", json={
        "completion_id": "cmp_damaged", "resource_path": str(path),
        "messages": [{"role": "user", "content": "你好"}],
    })
    assert response.status_code == 422
    assert resources[1] == before


@pytest.mark.parametrize("endpoint", ["/v1/document-processor/process", "/v1/document-processor/docx/process", "/v1/ocr/process"])
def test_old_processing_endpoints_removed(endpoint):
    response = TestClient(create_app()).post(endpoint, files={"file": ("a.pdf", b"pdf")})
    assert response.status_code == 404


def test_prepare_pdf_calls_parser_then_builds_resource(resources, monkeypatch):
    from service.document_processor import processor
    from service.document_processor.schemas import ProcessResult
    calls = []

    def parse(file, file_type=None):
        calls.append((file.filename, file.read()))
        return ProcessResult(filename=file.filename, html="<p>PDF 内容</p>")

    monkeypatch.setattr(processor, "process", parse)
    response = TestClient(create_app()).post("/v1/document-resources", files={"files": ("a.pdf", b"%PDF-1.4")})
    assert response.status_code == 200, response.text
    assert calls == [("a.pdf", b"%PDF-1.4")]
    assert response.json()["documents"] == [{"filename": "a.pdf", "html": "<p>PDF 内容</p>"}]
    assert resources[1] == ["PDF 内容"]


def test_parser_failure_identifies_file_and_does_not_build_index(resources, monkeypatch):
    from service.document_processor import processor

    def fail(file, file_type=None):
        raise RuntimeError("invalid PDF")

    monkeypatch.setattr(processor, "process", fail)
    response = TestClient(create_app()).post("/v1/document-resources", files={"files": ("bad.pdf", b"pdf")})
    assert response.status_code == 500
    assert "bad.pdf" in response.text and "invalid PDF" in response.text
    assert resources[1] == []
    assert list(resources[0].iterdir()) == []
