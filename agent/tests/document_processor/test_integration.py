from __future__ import annotations

from io import BytesIO

from docx import Document

from service.document_processor.processor import process


MINIMAL_PDF_BYTES = b"%PDF-1.4\n% TraceAgent synthetic PDF fixture\n%%EOF\n"


def build_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Overview", level=1)
    document.add_paragraph("Alpha paragraph.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_process_handles_pdf_file_path_via_public_interface(monkeypatch, tmp_path):
    from service.document_processor import processor as processor_module

    fixture_path = tmp_path / "sample_notice.pdf"
    fixture_path.write_bytes(MINIMAL_PDF_BYTES)
    seen_call: dict[str, object] = {}

    def fake_convert(source_bytes: bytes, filename: str) -> list[list[dict]]:
        seen_call["source_prefix"] = source_bytes[:4]
        seen_call["filename"] = filename
        return [[
            {
                "type": "paragraph",
                "content": {
                    "paragraph_content": [
                        {"type": "text", "content": "合成 PDF fixture"}
                    ]
                },
            }
        ]]

    monkeypatch.setattr(processor_module, "convert_pdf_bytes_to_content_list", fake_convert)

    with fixture_path.open("rb") as file_obj:
        result = process(file_obj)

    assert result.filename == fixture_path.name
    assert 'id="p001_b000"' in result.html
    assert "合成 PDF fixture" in result.html
    assert seen_call["source_prefix"] == b"%PDF"
    assert seen_call["filename"] == fixture_path.name


def test_process_routes_docx_file_path_via_public_interface(tmp_path):
    fixture_path = tmp_path / "sample_notice.docx"
    fixture_path.write_bytes(build_docx_bytes())

    with fixture_path.open("rb") as file_obj:
        result = process(file_obj)

    assert result.filename == fixture_path.name
    assert result.meta_info == {"engine": "python-docx"}
    assert result.blocks[0]["block_id"] == "docx_b001"
    assert result.blocks[0]["text"] == "Overview"
    assert '<h1 id="docx_b001"' in result.html
