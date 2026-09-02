from io import BytesIO

import pytest
from docx import Document

from service.document_processor.schemas import ProcessResult


class NamedBytesIO(BytesIO):
    def __init__(self, data: bytes, filename: str | None = None) -> None:
        super().__init__(data)
        self.filename = filename


def build_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Overview", level=1)
    document.add_paragraph("Alpha paragraph.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_process_validates_input_then_calls_pdf_pipeline(monkeypatch):
    from service.document_processor import processor as processor_module

    seen_call: dict[str, object] = {}

    def fake_convert(source_bytes: bytes, filename: str) -> str:
        seen_call["source_bytes"] = source_bytes
        seen_call["filename"] = filename
        return '<!doctype html><html><head><style>x</style></head><body><main><section class="page" id="page_001"><p id="p001_b000">正文</p></section></main></body></html>'

    monkeypatch.setattr(processor_module, "convert_pdf_to_html", fake_convert)

    file_obj = NamedBytesIO(b"%PDF-1.4", filename="/tmp/sample.PDF")
    result = processor_module.process(file_obj)

    assert result.filename == "sample.PDF"
    assert '<!doctype html>' in result.html
    assert "<style>" in result.html
    assert 'id="p001_b000"' in result.html
    assert "正文" in result.html
    assert seen_call["source_bytes"] == b"%PDF-1.4"
    assert seen_call["filename"] == "sample.PDF"
    assert file_obj.tell() == 0


def test_process_uses_mineru_even_when_pdf_text_layer_is_readable(monkeypatch):
    from service.document_processor import processor as processor_module

    seen_call: dict[str, object] = {}

    def fake_convert(source_bytes: bytes, filename: str) -> str:
        seen_call["source_bytes"] = source_bytes
        seen_call["filename"] = filename
        return "<p>MinerU正文</p>"

    monkeypatch.setattr(processor_module, "convert_pdf_to_html", fake_convert)

    result = processor_module.process(NamedBytesIO(b"%PDF-1.4", filename="text.pdf"))

    assert "MinerU正文" in result.html
    assert seen_call["source_bytes"] == b"%PDF-1.4"
    assert seen_call["filename"] == "text.pdf"


def test_process_accepts_explicit_pdf_type_without_filename_suffix(monkeypatch):
    from service.document_processor import processor as processor_module

    monkeypatch.setattr(
        processor_module,
        "convert_pdf_to_html",
        lambda source_bytes, filename: '<p id="p001_b000">正文</p>',
    )

    result = processor_module.process(
        NamedBytesIO(b"%PDF-1.4", filename="upload.bin"),
        file_type=".PDF",
    )

    assert result.filename == "upload.bin"
    assert 'id="p001_b000"' in result.html
    assert "正文" in result.html


def test_process_rejects_objects_without_file_like_read_method():
    from service.document_processor import processor as processor_module
    from service.document_processor.processor import InvalidFileObjectError

    with pytest.raises(InvalidFileObjectError, match="file-like"):
        processor_module.process(object(), file_type="pdf")


def test_process_routes_docx_explicit_type_to_docx_pipeline():
    from service.document_processor import processor as processor_module

    result = processor_module.process(
        NamedBytesIO(build_docx_bytes(), filename="sample.pdf"),
        file_type="docx",
    )

    assert result.filename == "sample.pdf"
    assert "Overview" in result.html
    assert "Alpha paragraph." in result.html


def test_process_routes_docx_filename_suffix_to_docx_pipeline():
    from service.document_processor import processor as processor_module

    result = processor_module.process(NamedBytesIO(build_docx_bytes(), filename="sample.DOCX"))

    assert result.filename == "sample.DOCX"
    assert "Overview" in result.html
    assert "Alpha paragraph." in result.html


def test_process_routes_docx_explicit_type_to_default_docx_filename():
    from service.document_processor import processor as processor_module

    result = processor_module.process(BytesIO(build_docx_bytes()), file_type="docx")

    assert result.filename == "document.docx"
    assert "Overview" in result.html


def test_process_rejects_unsupported_explicit_type():
    from service.document_processor import processor as processor_module
    from service.document_processor.processor import UnsupportedFileTypeError

    with pytest.raises(UnsupportedFileTypeError, match="txt"):
        processor_module.process(
            NamedBytesIO(b"fake", filename="sample.pdf"),
            file_type="txt",
        )


def test_process_rejects_non_pdf_filename_when_type_is_omitted():
    from service.document_processor import processor as processor_module
    from service.document_processor.processor import UnsupportedFileTypeError

    with pytest.raises(UnsupportedFileTypeError, match="txt"):
        processor_module.process(NamedBytesIO(b"fake", filename="sample.txt"))


def test_process_uses_default_pdf_filename_when_name_is_missing(monkeypatch):
    from service.document_processor import processor as processor_module

    seen_call: dict[str, object] = {}

    def fake_convert(source_bytes, filename):
        seen_call["filename"] = filename
        return "<style>body{}</style>正文"

    monkeypatch.setattr(
        processor_module,
        "convert_pdf_to_html",
        fake_convert,
    )

    result = processor_module.process(BytesIO(b"%PDF-1.4"))

    assert result.filename == "document.pdf"
    assert seen_call["filename"] == "document.pdf"
    assert "<style>" in result.html
