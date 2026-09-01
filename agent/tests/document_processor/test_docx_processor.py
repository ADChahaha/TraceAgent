from __future__ import annotations

from io import BytesIO

from docx import Document


class NamedBytesIO(BytesIO):
    def __init__(self, data: bytes, filename: str | None = None) -> None:
        super().__init__(data)
        self.filename = filename


def build_docx_bytes(*, with_headings: bool) -> bytes:
    document = Document()
    if with_headings:
        document.add_heading("Overview", level=1)
        document.add_paragraph("Alpha paragraph.")
        document.add_heading("Details", level=2)
        document.add_paragraph("Beta paragraph.")
    else:
        document.add_paragraph("Plain first paragraph.")
        document.add_paragraph("Plain second paragraph.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Term"
    table.cell(0, 1).text = "Notice"
    table.cell(1, 0).text = "Termination"
    table.cell(1, 1).text = "30 days"

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_process_docx_builds_sections_from_word_heading_styles():
    from service.document_processor import processor

    file_obj = NamedBytesIO(build_docx_bytes(with_headings=True), filename="/tmp/sample.DOCX")

    result = processor.process(file_obj)

    assert result.filename == "sample.DOCX"
    assert result.meta_info == {"engine": "python-docx"}
    assert result.warnings == []
    assert file_obj.tell() == 0
    assert result.blocks[0]["block_id"] == "docx_b001"
    assert result.blocks[0]["kind"] == "heading"
    assert result.blocks[0]["text"] == "Overview"
    assert result.blocks[0]["page_no"] is None
    assert result.blocks[1]["block_id"] == "docx_b002"
    assert result.blocks[1]["kind"] == "paragraph"
    assert result.blocks[1]["text"] == "Alpha paragraph."
    assert result.blocks[-1]["kind"] == "table"
    assert result.blocks[-1]["block_id"] == "docx_b005"
    assert "docx_b005_tr_002" in result.html
    assert 'id="docx_b001_section"' in result.html
    assert '<h1 id="docx_b001"' in result.html
    assert '<h2 id="docx_b003"' in result.html
    assert '<table id="docx_b005"' in result.html
    assert result.display_html is not None
    assert "<html" in result.display_html
    assert "docx_b005_tr_002" in result.display_html
    assert "# Overview" in result.markdown
    assert "## Details" in result.markdown
    assert "| Term | Notice |" in result.markdown
    assert result.md_list == [
        "Overview",
        "Alpha paragraph.",
        "Details",
        "Beta paragraph.",
        "Term Notice Termination 30 days",
    ]
    assert result.semantic_document["sections"] == [
        {
            "section_id": "docx_b001",
            "title": "Overview",
            "level": 1,
            "text": "Overview\nAlpha paragraph.\nDetails\nBeta paragraph.\nTerm Notice Termination 30 days",
            "block_ids": ["docx_b001", "docx_b002", "docx_b003", "docx_b004", "docx_b005"],
        },
        {
            "section_id": "docx_b003",
            "title": "Details",
            "level": 2,
            "text": "Details\nBeta paragraph.\nTerm Notice Termination 30 days",
            "block_ids": ["docx_b003", "docx_b004", "docx_b005"],
        },
    ]
    assert result.semantic_document["blocks"][-1]["block_id"] == "docx_b005"
    assert result.semantic_document["blocks"][-1]["rows"][1]["row_id"] == "docx_b005_tr_002"


def test_process_docx_without_heading_styles_keeps_flat_original_order():
    from service.document_processor import processor

    file_obj = NamedBytesIO(build_docx_bytes(with_headings=False), filename="flat.docx")

    result = processor.process(file_obj)

    assert [block["kind"] for block in result.blocks] == ["paragraph", "paragraph", "table"]
    assert [block["text"] for block in result.blocks] == [
        "Plain first paragraph.",
        "Plain second paragraph.",
        "Term Notice Termination 30 days",
    ]
    assert result.semantic_document["sections"] == []
    assert result.semantic_document["blocks"][0]["block_id"] == "docx_b001"
    assert '<section id=' not in result.html
    assert "<h1" not in result.html
    assert "Plain first paragraph." in result.markdown
    assert "| Term | Notice |" in result.markdown
